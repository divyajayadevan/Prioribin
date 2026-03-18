import subprocess
import sys
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Prioribin Project: Important Codes', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('This document contains the most important functions and code blocks for the Prioribin Project, organized by component.')
    
    snippets = [
        ("Function: calculate_status (app.py)", """def calculate_status(fill_level):
    if fill_level >= 90: return "Critical"
    elif fill_level >= 70: return "Warning"
    return "Normal\""""),
        ("Function: log_event (app.py)", """def log_event(bin_id, event_type, description, collector_name=None):
    new_log = BinHistory(
        bin_id=bin_id, 
        event_type=event_type, 
        description=description,
        collector_name=collector_name
    )
    db.session.add(new_log)"""),
        ("Function: admin_dashboard (app.py)", """@app.route('/admin')
def admin_dashboard():
    bins = WasteBin.query.order_by(WasteBin.fill_level.desc()).all()
    # Fetch active collectors for the new UI
    cutoff = datetime.utcnow() - timedelta(minutes=5)
    active_collectors = Collector.query.filter(Collector.last_active >= cutoff).all()
    all_collectors = Collector.query.all()
    return render_template('admin.html', bins=bins, active_collectors=active_collectors, all_collectors=all_collectors)"""),
        ("Function: bin_history (app.py)", """@app.route('/history/<bin_id>')
def bin_history(bin_id):
    logs = BinHistory.query.filter_by(bin_id=bin_id).order_by(BinHistory.timestamp.desc()).limit(50).all()
    
    import json
    import re
    graph_labels = []
    graph_data = []
    
    for log in reversed(logs):
        if log.event_type == 'Collection':
            graph_labels.append(log.timestamp.strftime('%m-%d %H:%M'))
            graph_data.append(0)
        elif log.event_type in ['Critical Alert', 'Update']:
            match = re.search(r'Sensor: (\d+)%', log.description)
            level = int(match.group(1)) if match else 100
            graph_labels.append(log.timestamp.strftime('%m-%d %H:%M'))
            graph_data.append(level)
            
    return render_template('history.html', logs=logs, bin_id=bin_id, chart_labels=json.dumps(graph_labels), chart_data=json.dumps(graph_data))"""),
        ("Function: collector_login (app.py)", """@app.route('/collector_login', methods=['GET', 'POST'])
def collector_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        collector = Collector.query.filter_by(username=username).first()
        if collector and collector.check_password(password):
            session['collector_id'] = collector.id
            session['collector_name'] = collector.name
            session['collector_username'] = collector.username
            collector.last_active = datetime.utcnow()
            db.session.commit()
            return redirect(url_for('collector_dashboard'))
        else:
            flash('Invalid username or password', 'error')
            
    return render_template('collector_login.html')"""),
        ("Function: collector_dashboard (app.py)", """@app.route('/collector')
def collector_dashboard():
    if 'collector_id' not in session:
        return redirect(url_for('collector_login'))
    
    collector_id = session['collector_id']
    collector = Collector.query.get(collector_id)
    if collector:
        collector.last_active = datetime.utcnow()
        db.session.commit()

    all_bins = WasteBin.query.all()
    return render_template('collector.html', bins=all_bins, collector_name=session.get('collector_name'), collector_username=session.get('collector_username'))"""),
        ("Function: register_collector (app.py)", """@app.route('/admin/register_collector', methods=['POST'])
def register_collector():
    name = request.form.get('name')
    username = request.form.get('username')
    password = request.form.get('password')
    
    if name and username and password:
        existing = Collector.query.filter_by(username=username).first()
        if not existing:
            new_collector = Collector(name=name, username=username)
            new_collector.set_password(password)
            db.session.add(new_collector)
            db.session.commit()
            flash(f'Collector {name} registered successfully!', 'success')
        else:
            flash('Username already exists.', 'error')
            
    return redirect(url_for('admin_dashboard'))"""),
        ("Function: update_location (app.py)", """@app.route('/api/update_location', methods=['POST'])
def update_location():
    \"\"\"Receives GPS from Collector Phone\"\"\"
    data = request.json
    username = data.get('username')  # Changed from name to username for unique lookup
    name = data.get('name') 
    lat = data.get('lat')
    lon = data.get('lon')

    if username:
        collector = Collector.query.filter_by(username=username).first()
    else:
        # Fallback for older apps temporarily
        collector = Collector.query.filter_by(name=name).first()

    if collector:
        if lat and lon: # only update if provided
            collector.lat = lat
            collector.lon = lon
        collector.last_active = datetime.utcnow()
        db.session.commit()
        return jsonify({"success": True})
    return jsonify({"error": "Collector not found"}), 404"""),
        ("Function: update_bin (app.py)", """@app.route('/api/update_bin', methods=['POST'])
def update_bin():
    \"\"\"Hardware Simulation Endpoint\"\"\"
    data = request.json
    bin_obj = WasteBin.query.filter_by(bin_id=data['bin_id']).first()
    if bin_obj:
        new_level = int(data['fill_level'])
        new_status = calculate_status(new_level)
        
        # Log event if level has changed to ensure graph history
        if bin_obj.fill_level != new_level:
            if bin_obj.status != "Critical" and new_status == "Critical":
                log_event(bin_obj.bin_id, "Critical Alert", f"Sensor: {new_level}%")
            else:
                # Log regular update
                log_event(bin_obj.bin_id, "Update", f"Sensor: {new_level}%")

        bin_obj.fill_level = new_level
        bin_obj.status = new_status
        db.session.commit()
        return jsonify({"status": bin_obj.status}), 200
    return jsonify({"error": "Bin not found"}), 404"""),
        ("Function: collect_bin (app.py)", """@app.route('/api/collect_bin/<bin_id>', methods=['POST'])
def collect_bin(bin_id):
    \"\"\"Collector marked bin as cleaned\"\"\"
    # Get collector name from JSON body
    data = request.json
    collector_name = data.get('collector_name', 'Unknown')
    
    bin_obj = WasteBin.query.filter_by(bin_id=bin_id).first()
    if bin_obj:
        log_event(bin_id, "Collection", "Cleaned by collector", collector_name)
        bin_obj.fill_level = 0
        bin_obj.status = "Normal"
        db.session.commit()
        return jsonify({"success": True}), 200
    return jsonify({"error": "Bin not found"}), 404"""),
        ("Class: WasteBin (models.py)", """class WasteBin(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    bin_id = db.Column(db.String(50), unique=True, nullable=False)
    location_lat = db.Column(db.Float, nullable=False)
    location_lon = db.Column(db.Float, nullable=False)
    fill_level = db.Column(db.Integer, default=0)
    status = db.Column(db.String(20), default="Normal")
    last_updated = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            "bin_id": self.bin_id,
            "lat": self.location_lat,
            "lon": self.location_lon,
            "fill_level": self.fill_level,
            "status": self.status
        }"""),
        ("Class: Collector (models.py)", """class Collector(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    lat = db.Column(db.Float, nullable=True)
    lon = db.Column(db.Float, nullable=True)
    last_active = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def to_dict(self):
        return {
            "name": self.name,
            "username": self.username,
            "lat": self.lat,
            "lon": self.lon,
            "last_active": self.last_active.isoformat()
        }"""),
        ("Function: get_registered_bins (simulate_hardware.py)", """def get_registered_bins():
    \"\"\"Fetches real bins from the Admin Dashboard\"\"\"
    try:
        response = requests.get(GET_BINS_URL)
        if response.status_code == 200:
            return response.json() # Returns list of bins
    except:
        return []
    return []"""),
        ("Function: main loop block (simulate_hardware.py)", """try:
    while True:
        # 1. Fetch Active Bins from Server
        # We check this every loop so if you add a new bin in Admin, it appears here automatically.
        active_bins = get_registered_bins()
        
        if not active_bins:
            print("⚠️  No bins found in system. Please add a bin in Admin Dashboard.")
            time.sleep(15)
            continue

        # Hardware bins! These are completely ignored by the simulator.
        # Any other bin ID you create in the Admin Dashboard will automatically get simulated.
        IGNORED_BINS = ["BIN-01", "BIN-02"]

        simulated_count = 0
        for bin_data in active_bins:
            b_id = bin_data['bin_id']
            
            # Skip any bin that is reserved for real physical hardware
            if b_id in IGNORED_BINS:
                continue
                
            simulated_count += 1
            server_fill_level = bin_data['fill_level'] # The level currently in DB
            
            # Initialize local state if new
            if b_id not in bin_states:
                bin_states[b_id] = server_fill_level

            # 2. LOGIC: Check if Collector emptied it
            # If server says 0 but we thought it was 100, the collector cleaned it!
            if server_fill_level == 0 and bin_states[b_id] > 0:
                print(f"♻️  [EVENT] {b_id} was emptied by Collector. Resetting Edge Sensor.")
                bin_states[b_id] = 0
            
            # 3. LOGIC: Simulate Waste Accumulation
            current_level = bin_states[b_id]

            if current_level >= 100:
                # Bin is full. It CANNOT go down unless collected.
                new_level = 100
                status_msg = "🚨 CRITICAL: OVERFLOW DETECTED"
            else:
                # Simulate people throwing trash (random increase 5% to 15%)
                increase = random.randint(5, 15)
                new_level = min(current_level + increase, 100) # Cap at 100
                status_msg = "✅ Monitoring... (Filling Up)"

            # Update local memory
            bin_states[b_id] = new_level

            # 4. OUTPUT (Matches PDF concept of Edge Processing)
            print("-" * 50)
            print(f"📦 BIN ID: {b_id}")
            print(f"   Sensor Reading: {new_level} cm (converted to %)")
            
            if new_level >= 90:
                print("   [EDGE LOGIC] 🛑 Threshold Exceeded -> PRIORITY HIGH")
                print("   STATUS: CRITICAL (Needs Collector)")
            elif new_level >= 70:
                print("   [EDGE LOGIC] ⚠️ Threshold Approaching -> PRIORITY MED")
            else:
                print(f"   [EDGE LOGIC] Normal accumulation (+{new_level - current_level if new_level>current_level else 0}%)")

            # 5. Send to Server (Only if changed or critical to keep heartbeat)
            payload = {"bin_id": b_id, "fill_level": new_level}
            try:
                requests.post(UPDATE_URL, json=payload, timeout=1)
            except:
                print("   ❌ Network Error: Server Offline")

            time.sleep(1) # Short pause between bins

        print("\\n⏳ Cycle complete. Waiting for next sensor reading...\\n")
        time.sleep(15) # Wait 15 seconds before next batch

except KeyboardInterrupt:
    print("\\n🛑 Simulation Stopped.")""")
    ]

    for title_text, code_text in snippets:
        doc.add_heading(title_text, level=2)
        
        # Add a paragraph for the code and change properties for each run
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        
        # Make the code font Courier/Monospace
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

        # Style the paragraph keeping lines separate
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.0

        # Adding a bit of shading or visual distinctiveness (optional)
        # For simple code representation, Courier New usually suffices in Word docs.

    # Save to file
    out_path = os.path.join("c:\\project_repos\\Prioribin_Project", "Important_Codes.docx")
    doc.save(out_path)
    print(f"Document saved to {out_path}")

if __name__ == '__main__':
    main()
